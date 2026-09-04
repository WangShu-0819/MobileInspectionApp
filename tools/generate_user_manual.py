from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"D:\study\Textile_defects\Wearable Inspection\MobileInspectionApp")
IMAGE_ROOT = Path(r"D:\study\刘老师\项目\AR眼镜\9.4 v1\图片")
OUTPUT = WORKSPACE / "docs" / "user-guide" / "视觉质检MobileInspectionApp_使用说明.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
MUTED = RGBColor(112, 122, 132)
GOLD = RGBColor(160, 111, 18)
RED = RGBColor(155, 28, 28)
GREEN = RGBColor(28, 125, 86)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color=RGBColor(32, 43, 54), bold=False):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    """Apply fixed DXA geometry to a real Word table."""
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_paragraph_border(paragraph, color="D7DBE2", size="8", space="6", side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_shaded_note(doc, label: str, text: str, color="F4F6F9", label_color=DARK_BLUE):
    p = doc.add_paragraph(style="Manual Note")
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)
    add_paragraph_border(p, color="9DB8CF", size="14", space="8", side="left")
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, color=label_color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=GRAY)
    return p


def add_body(doc, text: str, after=6, before=0, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, color=NAVY, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=RGBColor(32, 43, 54))
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=RGBColor(32, 43, 54))
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.8, color=RGBColor(32, 43, 54))
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.8, color=RGBColor(32, 43, 54))
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_image(doc, path: Path, alt: str):
    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    max_width = 5.85 if aspect >= 0.72 else 3.25
    max_height = 6.65
    width_in = min(max_width, max_height * aspect)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", alt)
            doc_pr.set("title", alt)
    return p


def add_figure_page(doc, number: int, title: str, filename: str, description: str, steps: Iterable[str], tip: str | None = None):
    add_heading(doc, f"{number}. {title}", 2)
    image_path = IMAGE_ROOT / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    add_image(doc, image_path, f"视觉质检界面：{title}")
    add_caption(doc, f"图 {number}  {title}（素材：{filename}）")
    add_body(doc, description)
    for step in steps:
        add_number(doc, step)
    if tip:
        add_shaded_note(doc, "提示", tip)
    doc.add_page_break()


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=RGBColor(32, 43, 54))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = doc.styles[list_style_name]
        set_style_font(style, size=10.8, color=RGBColor(32, 43, 54))
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Manual Note" not in [s.name for s in doc.styles]:
        note = doc.styles.add_style("Manual Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = doc.styles["Manual Note"]
    set_style_font(note, size=10.5, color=GRAY)
    note.paragraph_format.line_spacing = 1.2

    caption = doc.styles["Caption"]
    set_style_font(caption, size=9, color=MUTED)
    caption.paragraph_format.line_spacing = 1.1

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("视觉质检 · MobileInspectionApp 使用说明")
    set_run_font(hr, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("内部操作参考  |  第 ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)
    fr = fp.add_run(" 页")
    set_run_font(fr, size=9, color=MUTED)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("现场视觉质检")
    set_run_font(r, size=12, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("视觉质检 MobileInspectionApp")
    set_run_font(r, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("使用说明")
    set_run_font(r, size=22, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("模板配置 · DPM 绑定 · 多视角采集 · 人工确认 · 结果留存")
    set_run_font(r, size=13, color=GRAY)

    table = doc.add_table(rows=5, cols=2)
    set_table_geometry(table, [2160, 7200], indent_dxa=120)
    header = table.rows[0]
    mark_header_row(header)
    set_cell_shading(header.cells[0], "D6E4EF")
    set_cell_shading(header.cells[1], "D6E4EF")
    for cell, text in zip(header.cells, ("项目", "说明")):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    rows = [
        ("适用对象", "现场质检员、模板管理员、项目维护人员"),
        ("运行方式", "Android 手机离线使用，数据默认保存在本机"),
        ("编写依据", "MobileInspectionApp 当前工程代码 + 20 张操作截图"),
        ("版本口径", "以当前源码行为为准；截图用于定位按钮和理解流程"),
    ]
    for row, (label, value) in zip(table.rows[1:], rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(label)
        set_run_font(r0, size=10.5, color=DARK_BLUE, bold=True)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        set_run_font(r1, size=10.5, color=GRAY)

    add_shaded_note(
        doc,
        "阅读方式",
        "正文按实际业务操作顺序编排。素材文件夹中 3 张 ZIP 内容截图的文件写入时间早于手机操作截图，因此本手册将它们放在“检测报告已生成”之后的导出结果环节；20 张源图片均已纳入。",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("2026 年 9 月")
    set_run_font(r, size=11, color=MUTED)
    doc.add_page_break()


def add_quick_start(doc: Document):
    add_heading(doc, "一、先看懂完整流程", 1)
    add_body(doc, "视觉质检 App 的核心闭环是：先建立零件和模板，再按模板顺序采集现场照片；有 ROI 的视角进入人工确认页，完成每个 ROI 与整张照片的 OK/NG 选择后继续；全部视角完成后生成检测报告并导出 ZIP。")
    add_shaded_note(doc, "一页速用", "我的 → 模板配置 → 新建零件 → 拍摄/导入模板视角 → 每个视角编辑 ROI 并选择属性 → 现场采集 → 拍照 → ROI/总体人工确认 → 确认并继续 → 下载或分享 ZIP。", color="EAF3F8")
    add_heading(doc, "1.1 三个一级入口", 2)
    table = doc.add_table(rows=1, cols=3)
    headers = ["入口", "主要用途", "常用操作"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    for values in (
        ("现场采集", "选择零件、DPM 扫码、模板叠加、拍照", "切换视角、调透明度、进入人工确认"),
        ("追溯记录", "查看历史批次和今日统计", "筛选、选择、导出 ZIP、删除批次"),
        ("我的", "模板、零件、模板包、结果和设置", "配置基础数据、导入/导出模板包"),
    ):
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=10.2, color=RGBColor(32, 43, 54))
    set_table_geometry(table, [2500, 3350, 3510], indent_dxa=120)
    mark_header_row(table.rows[0])
    add_heading(doc, "1.2 使用前准备", 2)
    add_bullet(doc, "首次进入需要相机的页面时，按系统提示允许相机权限；若永久拒绝，请到系统设置重新开启。")
    add_bullet(doc, "准备好零件 ID、零件名称、可选型号，以及零件上的 DPM/Data Matrix 码。")
    add_bullet(doc, "模板可以用手机后置相机拍摄，也可以从相册导入；现场 DPM 识别只走手机相机实时扫描，不支持从相册导入码图。")
    add_bullet(doc, "确保手机有足够存储空间。原始照片、模板图片、ROI 确认记录和导出 ZIP 都可能占用本机空间。")
    add_heading(doc, "1.3 本版本必须知道的边界", 2)
    add_bullet(doc, "ROI 和总体 OK/NG 是人工确认结果；用户没有选择时不会默认 OK 或 NG，NG 结果仍会保存。")
    add_bullet(doc, "导出 CSV 中的软件检测结果保留为空（未执行），不能把“人工 OK”理解为算法 PASS。")
    add_bullet(doc, "“我的 → 检测结果”当前是结果管理入口的空状态；已完成批次的实际下载和分享入口是检测报告页或追溯记录卡片。")
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_quick_start(doc)

    add_heading(doc, "二、配置零件、模板和 ROI", 1)
    add_body(doc, "模板管理员通常只需首次配置，或在零件结构、拍摄角度变化时重新配置。配置完成后，现场人员可以直接从“现场采集”开始工作。")

    add_figure_page(
        doc, 1, "进入现场采集：尚未配置模板", "现场采集界面-模板配置.jpg",
        "启动后默认进入“现场采集”。页面上方显示当前零件选择框、扫一扫入口和 OCR 钢印入口；中部是实时相机画面；如果当前零件没有模板，页面会提示“暂无模板”，并提供“前往模板配置”。",
        [
            "首次使用或当前零件还没有视角时，点击“前往模板配置”；也可以从底部“我的”进入模板配置。",
            "如果相机尚未准备好，先等待相机初始化；相机不可用时不要连续点击拍照。",
        ],
        "现场采集页是日常作业入口，不建议在这里临时修改 ROI；模板、视角和 ROI 配置集中在“我的 → 模板配置”。",
    )
    add_figure_page(
        doc, 2, "模板配置列表", "配置模板-新建模板.jpg",
        "模板配置页按零件展示已有模板和视角数量。每个零件卡片显示视角数、DPM 绑定状态，右上角“+”用于创建新零件。点击零件卡片进入该零件的视角网格。",
        [
            "点击右上角“+”创建零件；如果已有零件，点击对应卡片进入详情。",
            "已绑定 DPM 的零件会显示 DPM 信息；未绑定时显示“未绑定 DPM”。",
        ],
        "零件 ID 是稳定标识，创建后应尽量保持不变；现场识别和模板包导入都依赖它。",
    )
    add_figure_page(
        doc, 3, "新建零件：填写 ID 和名称", "新建模板零件-填写ID名称.jpg",
        "新建零件对话框至少需要填写“零件 ID”和“零件名称”，型号和 DPM 码可以先留空，后续再配置。示例中零件 ID 为 02，名称为“实验”。",
        [
            "输入零件 ID：仅支持字母、数字、下划线和连字符，长度 1～64 位。",
            "输入零件名称；型号为可选项，便于现场区分同类零件。",
            "点击“创建/保存”完成创建；ID 重复或格式不合法时按提示修改。",
        ],
        "如果准备通过扫码绑定 DPM，创建时可以不填 DPM 码，创建完成后从详情页右上角二维码按钮绑定。",
    )
    add_figure_page(
        doc, 4, "零件详情：拍摄或导入模板视角", "零件创建完毕-配对DPM码-拍摄导入图片.jpg",
        "零件详情页显示当前零件的视角数量和 DPM 状态。尚无视角时，页面中央提示“暂无视角”；右下角提供两个操作：上方为从相册导入，下方为用手机相机拍摄新的模板视角。",
        [
            "如需绑定 DPM，点击右上角二维码图标进入扫码页面。",
            "点击右下角蓝色相机按钮，拍摄一个新的模板视角；拍摄成功后会回到详情页。",
            "点击右上角白色图片按钮，从相册多选模板图片，系统按导入顺序建立多个视角。",
        ],
        "模板参考图应尽量覆盖完整零件、光照均匀、角度稳定；模板图是后续透明叠加和 ROI 定位的依据。",
    )
    add_figure_page(
        doc, 5, "为零件绑定 DPM：实时扫一扫", "配对DPM码.jpg",
        "DPM 扫码页是全屏相机实时扫描界面。将零件上的 Data Matrix 码放入扫描框，识别成功后会显示解码文本；当前代码支持后置相机、闪光灯和工业 DPM 预处理链。",
        [
            "让 DPM 码尽量位于扫描框中央，保持手机稳定并避免强反光。",
            "光线不足时点击右上角闪光灯图标；识别成功后返回上一页。",
            "绑定时若该码已属于其他零件，系统会拒绝冲突绑定；先处理旧绑定再重试。",
        ],
        "现场采集顶部的“扫一扫”同样只支持手机相机实时扫码；未知 DPM 码会提示先在模板配置中绑定。",
    )
    add_figure_page(
        doc, 6, "模板视角完成：查看视角数量与 ROI 数量", "视角拍摄完毕.jpg",
        "零件详情页以网格显示模板视角。卡片中包含缩略图、视角名称和 ROI 数量；示例零件有 6 个视角，不同视角可以配置不同数量的 ROI。",
        [
            "点击某个视角卡片进入“模板详情”，查看图片、顺序、状态和 ROI 列表。",
            "如需增加视角，点击右下角蓝色相机按钮；如需从本机图片建立视角，点击上方图片按钮。",
            "长按视角卡片可删除该视角，系统会同时删除参考图片和已有 ROI 配置。",
        ],
        "视角顺序会影响现场采集顺序和导出目录；如现场有固定拍摄路线，请按路线建立或整理视角。",
    )
    add_figure_page(
        doc, 7, "进入模板详情：查看和编辑 ROI", "单个视角设置ROI.jpg",
        "模板详情页展示视角名称、所属零件、当前视角序号、启用状态、创建/更新时间、参考图片路径和 ROI 列表。示例中视角 1/6 已定义 1 个 ROI，属性显示为“部件”。",
        [
            "确认参考图片和视角序号无误；需要重新采集模板图时点击“重拍”。",
            "在“ROI 区域”卡片中查看 ROI 数量和每个 ROI 当前目标属性。",
            "点击“编辑 ROI”进入 ROI 编辑器，新增、移动、缩放、删除或修改目标属性。",
        ],
        "历史 ROI 如果没有目标属性，会显示“未选择”；不要根据 ROI 位置自行猜测属性，应在编辑器中补选。",
    )
    add_figure_page(
        doc, 8, "ROI 编辑器：选择目标属性", "单个视角ROI编辑-属性.jpg",
        "ROI 编辑器在模板图上叠加显示红色/绿色 ROI 框。底部显示已定义数量和“目标属性”，属性菜单提供三种稳定选项：螺纹、螺母、部件。",
        [
            "新建 ROI：点击“添加 ROI”，在模板图片的有效内容区域内拖拽绘制矩形。",
            "选择 ROI：点击或长按已有矩形；选中后可拖拽矩形移动，拖拽四角控制柄缩放。",
            "点击“选择”，从“螺纹 / 螺母 / 部件”中选择与该 ROI 对应的目标类型。",
            "需要修改已有 ROI 时，先选中目标 ROI，再点击“修改”重新选择属性。",
        ],
        "保存按钮在目标属性未选择时保持不可用；当前任务只配置属性和坐标，不执行自动检测。",
    )
    add_figure_page(
        doc, 9, "ROI 编辑器：保存完成", "ROI编辑完毕保存.jpg",
        "选择目标属性后，底部会显示具体属性（示例为“部件”），蓝色“保存”按钮可用。保存会同时保留 ROI 的位置和目标属性。",
        [
            "检查 ROI 框是否完整覆盖目标区域，并确认没有超出模板图片内容边界。",
            "确认“目标属性”正确后点击“保存”；不想修改时点击“取消”。",
            "返回模板详情，确认 ROI 列表中出现正确属性；重新进入页面仍应保持。",
        ],
        "同一个视角可以有多个 ROI，每个 ROI 可以选择不同属性；属性不是整个零件的全局属性。",
    )

    add_heading(doc, "三、现场采集：按视角拍照并人工确认", 1)
    add_body(doc, "模板配置完成后，日常操作都从“现场采集”开始。系统以当前零件和模板的视角顺序创建一个采集批次；每次拍照都关联到该批次、当前视角、模板和照片文件。")

    add_figure_page(
        doc, 10, "现场采集：模板透明叠加", "模板配置完毕开始采集.jpg",
        "配置完成后，现场采集页顶部显示当前零件（示例为“实验”），实时相机画面中叠加模板参考图和 ROI 引导框。下方显示当前视角序号、模板参考图、透明度滑块和“拍照”按钮。",
        [
            "在顶部零件选择框中确认当前零件；也可以点击“扫一扫”扫描已绑定 DPM，自动切换零件及其模板。",
            "用模板参考图辅助摆放现场零件，观察实时画面和下方参考图是否一致。",
            "拖动“透明度”滑块调节模板叠加强度；点击眼睛图标可临时隐藏/显示模板。",
            "画面稳定、目标清晰后点击“拍照”，等待保存完成。",
        ],
        "当前页的模板叠加用于人工对位和拍摄辅助，不代表已经完成自动轮廓对齐或自动检测。",
    )
    add_figure_page(
        doc, 11, "切换现场采集视角", "切换采集视角.jpg",
        "点击视角选择区域后弹出“切换视角”底部面板。面板列出模板视角的顺序和名称，当前视角带有绿色勾选。",
        [
            "点击目标视角名称切换模板参考图、视角序号和 ROI 配置。",
            "切换前确认上一视角的照片和人工确认已经完成；未完成的视角不要用切换来代替确认。",
            "按模板顺序工作时，优先从 1/6 依次拍到 6/6，避免漏拍。",
        ],
        "每张现场照片按实际拍摄视角写入当前批次；不要只按列表位置判断照片归属。",
    )
    add_figure_page(
        doc, 12, "拍照完成后：逐个 ROI 和总体结果人工确认", "单个视角采集完毕人工勾选RO区域检测.jpg",
        "有 ROI 的视角拍照成功后会进入人工确认页。每张 ROI 卡片显示 ROI 子图、ROI 名称、ID、目标属性和 OK/NG；底部单独提供整张照片的“总体结果”OK/NG。",
        [
            "逐张查看 ROI 子图，点击该 ROI 的“OK”或“NG”；每个 ROI 都必须选择一次。",
            "在底部“总体结果”中独立选择整张照片的 OK 或 NG；总体结果不会由 ROI 结果自动计算。",
            "确认所有 ROI 和总体结果后，“确认并继续”按钮才会可用；点击后保存本视角确认记录并进入下一视角。",
            "如果任意 ROI 或总体结果为 NG，仍点击“确认并继续”完成保存；NG 不会被丢弃。",
        ],
        "未选择时不会默认 OK/NG。当前源码把软件检测结果保留为 null/未执行，人工确认仅表示现场人员的复核结论。无 ROI 的视角仍会保存照片并直接进入下一视角。",
    )

    add_heading(doc, "四、采集完成：生成报告并导出结果", 1)
    add_body(doc, "当批次的所有模板视角都完成拍照后，系统结束采集批次并进入检测结果页。只有完成全部视角并且照片索引完整时，批次才允许导出 ZIP。")
    add_figure_page(
        doc, 13, "所有视角完成：检测报告已生成", "所有视角采集完毕-生成检测报告-下载分享.jpg",
        "结果页显示“检测报告已生成”，并汇总零件名称、照片数量、检测记录数量和批次 ID。页面提供“下载 ZIP”“分享”和“返回采集页”。",
        [
            "先核对零件名称、照片数量和批次 ID，确认这是本次采集。",
            "点击“下载 ZIP”，在系统文件选择器中选择保存位置；下载成功后可在文件管理器中打开。",
            "点击“分享”可调用系统分享面板，把生成的 ZIP 交给协作软件或文件传输工具。",
            "完成交付后点击“返回采集页”，开始下一批次或切换零件。",
        ],
        "采集未完成时不会生成可导出的结果包；追溯记录中的未完成批次也会显示“拍完全部视角后才能导出 ZIP”。",
    )
    add_figure_page(
        doc, 14, "ZIP 根目录：照片目录和综合 CSV", "zip包.png",
        "检测结果 ZIP 的根目录包含 views 文件夹和 inspection_result.csv。views 保存本批次照片，CSV 是 Excel/WPS 可打开的综合索引和确认记录表。",
        [
            "打开 ZIP 后先确认存在 views 文件夹和 inspection_result.csv。",
            "如果只需要核对照片，进入 views；如果需要核对 ROI 结果、时间和关联关系，打开 CSV。",
            "不要把其他批次或其他零件的照片手工放入同一个 ZIP，以免破坏批次边界。",
        ],
        "CSV 是 Excel 兼容的 CSV，不是独立的 .xlsx 文件；用 Excel/WPS 打开时如遇中文显示异常，选择 UTF-8 编码导入。",
    )
    add_figure_page(
        doc, 15, "ZIP 内照片：按视角分目录", "zip包-图片.png",
        "views 文件夹按视角拆分为 view_01、view_02 等子目录。系统按实际 View 写入照片；同一视角重拍的照片也会保留，并通过唯一文件名避免覆盖。",
        [
            "按 view_01、view_02 等目录检查每个视角是否有照片。",
            "重点检查数量是否与结果页的照片数量一致；如果有重拍，注意同一视角可能有多张照片。",
            "查看照片时以 ZIP 内目录和 CSV 的照片索引为准，不要仅按文件名推断视角。",
        ],
        "无 ROI 的视角没有 ROI 确认行，但原始照片仍会进入 ZIP；这是正常行为。",
    )
    add_figure_page(
        doc, 16, "inspection_result.csv：照片索引与 ROI 确认", "zip包-excel.png",
        "综合 CSV 同时包含照片行和 ROI 确认行。表头覆盖零件、模板、View、ROI、ROI 属性、归一化坐标、现场像素坐标、人工确认结果、确认时间、总体结果和 ZIP 内照片路径等信息。",
        [
            "先看记录类型：照片行用于照片索引，ROI确认行用于逐个 ROI 的人工确认。",
            "查看“ROI属性”确认模板配置的螺纹/螺母/部件没有串用；查看“人工确认结果”确认 OK/NG。",
            "检查“总体结果”和“总体确认时间”，它们是整张照片的独立人工结论。",
            "如果表格中出现 ####，通常是列宽不足；在 Excel/WPS 中拉宽日期列即可，不代表数据为空。",
        ],
        "软件检测结果字段可能为空，表示本版本未执行自动 Detector；不要用空值替换或手工改成 PASS/FAIL。",
    )

    add_heading(doc, "五、追溯记录：筛选、导出和清理采集批次", 1)
    add_body(doc, "追溯记录保存每次现场采集批次。批次以稳定 batchId 关联照片和人工确认记录；卡片选择、删除和导出都针对 batchId，不按列表位置或零件名称误删。")
    add_figure_page(
        doc, 17, "追溯记录：筛选、选择、导出和删除", "追溯采集记录-删除-下载.jpg",
        "追溯记录页顶部显示今日统计：通过、不通过、待复核。下方“采集批次”可切换今日、近 3 天、近 7 天和所有；每张批次卡片显示零件、采集时间、视角数、批次 ID、选择框和导出按钮。",
        [
            "点击筛选器选择时间范围；默认显示“近 7 天”。",
            "点击批次卡片或右侧复选框选中批次，选中后卡片会突出显示，顶部垃圾桶按钮可用。",
            "点击“导出 ZIP”导出已完成批次；采集中的批次会禁用导出按钮。",
            "点击垃圾桶后核对弹窗中的批次数量和批次信息，确认后删除所选批次的照片和确认记录。",
        ],
        "批次删除不可恢复，但不会删除模板配置；如果只是想移除模板而保留历史照片，应使用模板包删除或模板/视角管理入口。",
    )

    add_heading(doc, "六、模板包和零件管理", 1)
    add_body(doc, "模板包适合在设备之间迁移或备份模板配置；零件管理适合维护零件主数据。两者的删除范围不同，操作前请先确认目标。")
    add_figure_page(
        doc, 18, "模板包：导入、导出和删除", "模板包的导入和导出.jpg",
        "模板包页面提供“选择 ZIP 文件”导入入口，并列出已配置模板包。每个模板包卡片显示零件、ID、视角数量和可选 DPM，下面提供“导出”和“删除”。",
        [
            "导入：点击“选择 ZIP 文件”，选择符合格式的模板 ZIP；系统会校验零件、视角图片、顺序和 ROI 配置。",
            "导出：在对应零件卡片点击“导出”，再在系统文件选择器中选择保存位置。",
            "导入成功后检查模板数量、ROI 数量和警告信息；缺图或非法 ROI 可能被跳过并提示。",
            "删除模板包会删除该零件的模板视角、参考图片和 ROI，但会保留历史采集批次和照片。",
        ],
        "模板包不是现场检测结果 ZIP：模板包用于配置迁移，检测结果 ZIP 用于交付现场照片和确认记录。",
    )
    add_figure_page(
        doc, 19, "零件管理：左滑显示删除", "零件管理-左滑删除.jpg",
        "零件管理页显示所有零件及其型号、DPM 信息。对某个零件卡片从右向左滑动，会露出红色“删除”操作。右上角“+”可新建零件。",
        [
            "进入“我的 → 零件管理”，找到需要维护的零件。",
            "从零件卡片右侧向左滑动，显示红色删除区域；松手后进入二次确认。",
            "如需新建零件，点击右上角“+”，填写 ID、名称，可选填型号和 DPM 码。",
        ],
        "零件删除会连带删除该零件下的模板和视角；历史检测记录按当前实现保留。",
    )
    add_figure_page(
        doc, 20, "删除零件：确认删除范围", "零件管理-删除零件确认.jpg",
        "删除确认框会明确提示零件名称和删除范围：该零件下的模板和视角会一起删除，历史检测记录会保留。只有点击“删除”后才执行。",
        [
            "核对弹窗中的零件名称，避免把相邻零件误删。",
            "需要保留配置时点击“取消”；确定不再使用该零件时点击“删除”。",
            "删除完成后返回零件列表；如需恢复，只能重新创建零件并重新配置或导入模板包。",
        ],
        "零件删除属于高影响操作。若只是清理某个模板包，不要从这里删除整个零件。",
    )

    add_heading(doc, "七、设置与常见问题", 1)
    add_heading(doc, "7.1 应用设置", 2)
    add_body(doc, "进入“我的 → 应用设置”可调整当前代码已接入的两类设置：DPM 网格重建尺寸和相机预览显示比例。")
    add_bullet(doc, "DPM 网格重建尺寸：自动（16/18/20）、16×16、18×18、20×20。默认自动；固定尺寸只适用于已知码制场景。")
    add_bullet(doc, "预览显示比例：原比例会完整显示相机画面，可能保留黑边；填充预览会铺满区域，但边缘可能被裁切。设置只影响预览显示，不改变拍照文件和坐标数据。")
    add_bullet(doc, "版本与诊断区域可查看当前应用版本、构建类型和包名。")

    add_heading(doc, "7.2 常见问题处理", 2)
    add_shaded_note(doc, "相机未就绪", "先等待初始化完成；仍失败时点击重试。若提示相机权限被拒绝，到系统设置开启相机权限；若提示相机被其他应用占用，关闭其他相机应用后再试。", color="FFF8E8", label_color=GOLD)
    add_shaded_note(doc, "没有模板", "返回“我的 → 模板配置”，选择零件后拍摄或导入至少一个模板视角；配置完成后回到现场采集。", color="F4F6F9")
    add_shaded_note(doc, "DPM 未找到", "先确认码面清晰、无强反光、位于扫描框附近；未知码需要先在模板配置中绑定到零件。现场扫码不支持相册导入码图。", color="F4F6F9")
    add_shaded_note(doc, "ROI 无法保存", "检查是否选中了 ROI，并且每个新建或编辑的 ROI 都已选择螺纹、螺母或部件；“未选择”不会自动猜测。", color="F4F6F9")
    add_shaded_note(doc, "ZIP 无法导出", "先确认该批次所有视角都有照片且采集已结束；采集中的批次会禁用导出。若仍失败，检查照片文件是否存在、手机存储空间和保存位置权限。", color="F4F6F9")
    add_shaded_note(doc, "结果怎么看", "本版本以人工确认结果为准：ROI 和总体 OK/NG 会保存到数据库并写入 CSV；软件检测结果保持空值/未执行，不能把人工 OK 当成算法 PASS。", color="EAF3F8")

    add_heading(doc, "八、操作完成检查清单", 1)
    add_body(doc, "交接或提交结果前，建议按以下顺序快速核对：")
    for item in (
        "当前零件名称、零件 ID 和 DPM 绑定正确。",
        "模板视角数量和采集计划一致，视角顺序没有遗漏。",
        "每个启用 ROI 都有正确的目标属性：螺纹、螺母或部件。",
        "每个视角都完成拍照；有 ROI 的视角完成全部 ROI 和总体 OK/NG 确认。",
        "结果页显示报告已生成，照片数量和批次 ID 与本次任务一致。",
        "ZIP 根目录包含 views 和 inspection_result.csv；views 内各视角照片齐全。",
        "CSV 中人工确认结果、总体结果和时间字段可追溯，NG 行没有被删除。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "附录：当前版本数据与功能边界", 1)
    add_body(doc, "当前工程把零件、模板、ROI、采集批次、照片和人工确认记录分开保存，并通过 batchId、photoId、templateId、viewIndex 和 roiId 关联。这样可以在同一零件的多个视角之间保持数据隔离，也能在导出 CSV 中还原“哪张照片、哪个视角、哪个 ROI、谁确认、何时确认”。")
    add_body(doc, "当前源码已接入模板透明叠加、DPM 实时扫码、钢印 OCR 入口、模板包导入/导出、现场多视角拍照、人工确认、批次 ZIP/CSV 导出和追溯批次管理。实时轮廓投影、自动对齐、自动 Detector/PASS-FAIL、结果管理列表等属于后续扩展边界，使用本说明时不要把界面预留入口理解为已经完成自动检测。")
    add_shaded_note(doc, "数据安全", "应用页面显示“离线可用”和“数据仅保存在本机”。导出或分享 ZIP 后，文件将受系统文件管理和分享渠道管理；交付后请按项目要求保管或清理导出文件。", color="FFF8E8", label_color=GOLD)

    # Remove the final empty page-break paragraph if present.
    if doc.paragraphs and doc.paragraphs[-1].text == "":
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
